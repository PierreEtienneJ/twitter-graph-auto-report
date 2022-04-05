from docx import Document
from docx.shared import Inches
import docx

import numpy as np
import pandas as pd
import collections
import matplotlib.pyplot as plt
import textdistance
import statistics 
import igraph

import os

import time
import json
import datetime

def dateTwitter2Timestamp(date):
    """Take date with format ddd mmm dd HH:MM:SS +0000 yyyy 
    and return 
    """
    month={"Jan":"01", "Feb":"02","Mar":"03", "Apr":"04", "May":"05", "Jun":"06", "Jul":"07", "Aug":"08", "Sep":"09", "Oct":"10", "Nov":"11", "Dec":"12"}
    return datetime.datetime.fromisoformat(date[-4:]+"-"+month[date[4:7]]+"-"+date[8:10]+"T"+date[11:19]).timestamp()

def processSentence(sentence:str)->str:
    """
    Deletion of some characters in the sentence

    Del emoji, and special characters, change flag emoji to country name.

    :param str sentence: input sentence
    :returns: new sentence without some special characters
    :rtype: string
    """
    caraToDel="🐝🚲⛺🌻🌲⭐⚽👨🏼‍🏫🕵🧭🇯🇻🏼‍😃👨🏻‍🏫🌎📲🏛🗽🌊🚢"  #del som emoji
    sentence=sentence.lower()
    transTable = sentence.maketrans("éèà-'.ôç≠ę:î", "eea   oc e i", "!#$%^&*()\"’‘«»,，@_/|"+caraToDel)
    sentence = sentence.translate(transTable)
    #translate some emoji
    for drap, txt in zip(["🇫🇷", "🇨🇳", "🇹🇼","🇰🇷","🇰🇵","🇩🇪","🇺🇲","🇪🇺", "🇷🇺", "🇮🇷","🇺🇳"], ["france","china", "taiwan", "south korea", "north korea", "germany", "usa", "europe", "russia", "iran","un"]):
        if(drap in sentence):
            sentence=sentence[:sentence.index(drap)]+" "+txt+" "+sentence[sentence.index(drap)+2:]
    while "\n" in sentence:
        sentence=sentence[:sentence.index("\n")]+" "+sentence[sentence.index("\n")+1:]
    while "  " in sentence:
        sentence=sentence[:sentence.index("  ")]+" "+sentence[sentence.index("  ")+2:]
    return sentence

def processListSentences(L:list):
    """
    Change a list of input sentences to a list clean sentences with :func:`processSentence` fuction

    :param list<str> L: Sentences list

    :returns: clean sentences
    :rtype: list
    """
    L_phrase=[]
    for l in L:
        a=processSentence(str(l))
        L_phrase.append(a)
    return L_phrase

def most_commun_word(L:list, min_word_size:int=3, process_word:bool=False):
    """
    :param list<str> L: Sentences list
    :param int min_word_size: minimum size of the words to be taken into consideration
    :param bool process_word: if the sentence has already been processed in :func:`processSentence`

    :returns: Return most commun word form a sentences list
    :rtype: dict
    """
    L_mot=[]
    for l in L:
        if(not process_word):
            l=processSentence(l)
        words= l.split(" ") #split each words
        for word in words:
            if(len(word)>=min_word_size): #dont keep small words
                L_mot.append(word)
    
    return collections.Counter(L_mot) #return dict {word:number of repeat}

def similarityBetween(target_account_description:str, follower_description:str, type_="hamming")->float:
    """
    Function to calcul similarity between two sentences (follower's description).
    This methode calcul similarity between two sentences taking into consideration the characters and not the meaning of the sentence

    :param str target_account_description: the first description (cleaning sentence)
    :param str follower_description: the second description (cleaning sentence)
    :param str type_: hamming or levenshtein

    :returns: similarity between two sentences
    :rtype: float
    """
    if(type_=="hamming"):
        return textdistance.hamming.similarity(target_account_description, follower_description) #distance #levenshtein #mlipns
    elif(type_=="levenshtein"):
        return textdistance.levenshtein.similarity(target_account_description, follower_description) #distance #levenshtein #mlipns

def writeDoc(file_name, file_content, previusDoc=False, deg_title=0):
    """title and text_1 are required, img and text_2 not. text_* can be a str, a dict {"text":'', "type":["bold", "italic"]} or a list of string to make bullet list. 

    :param str file_name: name of the doc (.docx)
    :param list file_content: [{"title":"",
                                "text_1":"",
                                "img":"",
                                "text_2":""},
                                [{"title":...}]]
    :param bool previusDoc: if previusDoc is True, the fonction continue the document, always False for user
    :param int deg_title: always 0 for user
    """
    
    if previusDoc:
        document = docx.Document(file_name)
    else:
        document = docx.Document()
    document.save(file_name)
    for content in file_content:
        if(type(content)==list):
            writeDoc(file_name, content, previusDoc=True, deg_title=deg_title+1)
            
        else: #(type(content)==dict):
            document = docx.Document(file_name)

            document.add_heading(content["title"], deg_title)
            
            if(type(content["text_1"])==dict):
                p = document.add_paragraph("")
                if("italic" in content["text_1"]["type"]):
                    p.add_run(content["text_1"]["text"]).italic = True
                if("bold" in content["text_1"]["type"]):
                    p.add_run(content["text_1"]["text"]).bold = True
            elif(type(content["text_1"])==list):
                for t in content["text_1"]:
                    p=document.add_paragraph(t, style='List Bullet')

            else:
                p = document.add_paragraph(content["text_1"])
    
            if("img" in content):
                document.add_picture(content["img"], width=Inches(5))
                
            if("text_2" in content):
                if(type(content["text_2"])==dict):
                    p = document.add_paragraph("")
                    if("italic" in content["text_2"]["type"]):
                        p.add_run(content["text_2"]["text"]).italic = True
                    if("bold" in content["text_2"]["type"]):
                        p.add_run(content["text_2"]["text"]).bold = True
                elif(type(content["text_1"])==list):
                    for t in content["text_1"]:
                        p=document.add_paragraph(t, style='List Bullet')
                else:
                    p = document.add_paragraph(content["text_2"])

            document.save(file_name)
        
class autoReport:
    """Class to create an analyse docx after graph analyse : neuds distributed in class and size.

    Use only the constructor and the methode makeReport
    """
    def __init__(self, nodes_csv:str, edges_csv:str, classes_param:dict, account:dict, **kwargs):
        """
        :param str nodes_csv: path and nodes file
        :param str edges_csv: path and edges file
        :param dict classes_param: {class ID:color} color compatible with matplotlib : https://matplotlib.org/stable/gallery/color/named_colors.html
        :param dict account: {name:<str>, at:<str>, description:<str>}
        
        :param str tweets_path: tweets path of a json (kwargs default None)
        
        :param str desc_col_name: name of the description column nodes.csv (kwargs default "description")
        :param dict id_col_name: name of the id column nodes.csv (kwargs default "Id")
        :param str class_col_name: name of the class column nodes.csv (kwargs default "modularity_class")
        :param dict rank_col_name: name of the rank column nodes.csv (kwargs default "pageranks")
        :param str name_col_name: name of the name account column nodes.csv (kwargs default "name")
        :param dict at_col_name: name of the @ account column nodes.csv (kwargs default "screen_name")
        :param str source_col_name: name of the source column in edges.csv (kwargs default "Source")
        :param dict target_col_name: name of the target column in edges.csv (kwargs default "Target")
        :param str path_out: name of the out folder (kwargs default temp)
        """
        self.target_account=account
        self.classes=classes_param
        self.file_nodes_csv=nodes_csv
        self.file_edges_csv=edges_csv
        
        self._desc_col_name=kwargs.get("desc_col_name","description")
        self._id_col_name= kwargs.get("id_col_name","Id")
        self._class_col_name= kwargs.get("class_col_name","modularity_class")
        self._rank_col_name= kwargs.get("rank_col_name","pageranks")
        self._name_col_name= kwargs.get("name_col_name","name")
        self._at_col_name= kwargs.get("at_col_name","screen_name")
        
        self._source_col_name= kwargs.get("source_col_name","Source")
        self._target_col_name=kwargs.get("target_col_name","Target")
        
        self.node=pd.read_csv(self.file_nodes_csv).loc[:, [self._id_col_name, self._name_col_name,self._at_col_name, self._desc_col_name, self._class_col_name,self._rank_col_name]]
        self.edge=pd.read_csv(self.file_edges_csv).loc[:, [self._source_col_name, self._target_col_name]]
        self.nb_nodes= len(self.node)
        
        # out path 
        self.out_path=kwargs.get("path_out","./temp/")
        if(self.out_path[-1]!='/'):
            self.out_path+='/'
        if not os.path.exists(self.out_path):
            os.makedirs(self.out_path)
            
        #vérifie les classes
        verif={key:False for key in self.classes}
        i=0
        while sum([v for _,v in verif.items()])!=len(verif) and i<len(self.node)-1:
            class_=self.node.loc[i, self._class_col_name]
            if(class_ in verif):
                verif[self.node.loc[i, self._class_col_name]]=True
            i+=1
            
        if(sum([v for _,v in verif.items()])!=len(verif)):
            raise Exception(f"classes {[v for v,a in verif.items() if not a]} have no nodes")
        
        #tweet to dataframe
        self._tweets=kwargs.get("tweets_path",None)
        self._df_tweets=None
        
        if self._tweets:
            L=[]
            for t in json.load(open(self._tweets, "r")):

                inter = "response" if "in_reply_to_screen_name" in t else "tweet"
                if("quoted_status" in t):
                    inter ="rt" 

                L.append({
                    "created_at":datetime.datetime.fromtimestamp(dateTwitter2Timestamp(t["created_at"])).isoformat(),
                    "text": t["full_text"] if "full_text" in t else t["text"],
                    "user_screen_name":t["user"]["screen_name"],
                    "user_id":t["user"]["id_str"],
                    self._class_col_name:self.node[self.node["Id"] == t["user"]["id"]]["modularity_class"].tolist()[0] if not self.node[self.node["Id"] == t["user"]["id"]].empty else None,
                    "interaction":inter,
                    "language":t["metadata"]["iso_language_code"]
                })
            self._df_tweets=pd.DataFrame(L)
    
    def pieGraph(self,file_name, datas, title):
        """
        """
        plt.clf() 
        data=[v/sum(list(datas.values())) for v in datas.values()]
        nData=[]
        nLabel=[]
        if(len(datas)>3):
            other=0
            for d,lab in zip(data, list(datas.keys())):
                if(d<0.01):
                    other+=d
                else:
                    nData.append(d)
                    nLabel.append(lab)
            if(other!=0):
                nData.append(other)
                nLabel.append("other")

        else:
            nData=data
            nLabel=list(datas.keys())

        plt.pie(nData, labels = nLabel) 
        plt.title(title)
        plt.savefig(file_name, bbox_inches = 'tight')

    def class_distribution(self, file_name:str)->list:
        """
        Create a image file with a graph bar to show class distribution

        :param str file_name: name of the saving image (.png, .jpg...)
        :returns: list of distribution
        :rtype: list<float>
        """
        plt.clf()
        class_=collections.Counter(self.node.loc[:, self._class_col_name].tolist())
        
        plt.bar(list(self.classes.values()), [class_[c]/self.nb_nodes for c in self.classes.keys()], color=list(self.classes.values()))
        plt.xticks(rotation=45)
        plt.savefig(file_name, bbox_inches = 'tight')
        
        return [class_[c]/self.nb_nodes for c in self.classes.keys()] #return percentage 
    
    def most_common_words_tot(self, file_name:str, nb_words:int, min_len_word:int)->None:
        """
        Create a graph file with all most common words on followers description regardless of their class

        :param str file_name: name of the saving image (.png, .jpg...)
        :param int nb_words: number of the most common words to be displayed on the graph
        :param int min_len_word: minimum size of the words to be taken into consideration look :func:`most_commun_word`
        """
        plt.clf() #clear graph
        fig, ax = plt.subplots()

        countL_tot=most_commun_word(processListSentences(self.node.loc[:,self._desc_col_name].tolist()), min_len_word, True) #take all most commun word
        
        if("https" in countL_tot): #del https, it's not intresting
            countL_tot.pop("https")
        
        countL_tot=countL_tot.most_common(nb_words) #keep just some most common words
        
        #creation of the bar graph
        X = np.arange(len(countL_tot))*1.5
        for i,key in enumerate(self.classes.keys()):
            rslt_df = self.node[self.node[self._class_col_name] == key]
            L_phrase=processListSentences(rslt_df.loc[:,self._desc_col_name].tolist())
            countL=most_commun_word(L_phrase, min_len_word, True)

            ax.bar(X + (i/len(self.classes)-0.5), [countL[k]/val for k,val in countL_tot], color=self.classes[key], width = 1/len(self.classes))
        ax.set_xticks(X)
        ax.set_xticklabels([word for word, _ in countL_tot], rotation=45)
        ax.set_ylabel(f"distribution of the {len(countL_tot)} most common words by class")
        plt.savefig(file_name, bbox_inches = 'tight')
        plt.clf() #clear graph
    
    def most_commun_words_by_class(self, file_name:str, sentences_list:list, nb_words:int, min_len_word:int, title:str=None, color:str="blue", **kwargs):
        """Create a bar graph with a words list for a single class; use for classes descriptions and tweets
        
        :param str file_name: name of the saving image (.png, .jpg...)
        :param list<str> sentences_list: list for sentences
        :param int nb_words: number of the most common words to be displayed on the graph
        :param int min_len_word: minimum size of the words to be taken into consideration look :func:`most_commun_word`
        :param str color: graph color default blue
        :param list<str> kwargs/words_notIncluded: words not included in the graph default ['https']
        :param str kwargs/title: graph title default f"The {len(L_)} most common word in {color} class"
        """
        plt.clf() #clear graph
        #rslt_df = self.node[self.node[self._class_col_name] == key_class]
        countL=most_commun_word(processListSentences(sentences_list), min_len_word, True)
        for w in kwargs.get("words_notIncluded", ['https']):
            if(w in countL):
                countL.pop(w)
            
        L_=countL.most_common(nb_words)
        
        plt.bar([k for k, _ in L_], [val/len(sentences_list) for _, val in L_], width=1, color=color) #sum([val for _,val in L_])
        plt.xticks(rotation=45)

        #  {sum([val for k,val in class_ if k==key])/sum([val for _,val in class_])*100 :0.1f}%
        plt.title(kwargs.get("title", f"The {len(L_)} most common word in {color} class"))
        plt.savefig(file_name, bbox_inches = 'tight')
        
    def similarity_descriptions(self, file_name:str, similarity_fct):
        """
        Create a boxplot graph with similarity/distance between account target description and follower description; depending of their class

        :param str file_name: name of the saving image (.png, .jpg...)
        :param fct similarity_fct: function(str,str)->float, must return a similarity/distance 
        """
        plt.clf()
        result={}
        for key in self.classes.keys():
            rslt_df = self.node[self.node[self._class_col_name] == key]
            L_sentences=processListSentences(rslt_df.loc[:,self._desc_col_name].tolist())
            L_similarity=[]
            for sentence in L_sentences:
                L_similarity.append(similarity_fct(self.target_account["description"], sentence))

            result[key]=L_similarity #statistics.quantiles(data, *, n=4

        fig, ax = plt.subplots()
        ax.set_title(f"Similarities between the descriptions of the individuals in the class and {self.target_account['name']}'s description")
        ax.boxplot([result[key] for key in self.classes.keys()], showfliers=True, showmeans=True)
        ax.set_xticks(range(1,len(self.classes)+1))
        ax.set_xticklabels(list(self.classes.values()))
        ax.set_xlabel(f"class")
        plt.savefig(file_name, bbox_inches = 'tight')
    
    def links_between_class(self, file_name:str, percentage:float=1,node_weigth:list=None,logs:bool=True, **kwargs):
        """
        Create a node graph to resume the global graph, show all edges between classes
        This function can take time !

        :param str file_name: name of the saving image (.png, .jpg...)
        :param float percentage: must be in ]0,1], percentage of edge taken into consideration (in ordre of edges.csv) (default 1)
        :param list<float> node_weigth: same length as number of classes, weight of nodes, If None all nodes have the same weight (default None)
        :param bool logs: if plot some logs to show the progress (default True)
        :param bool show_internal_link: if true, draw an edge from the node to the same node (default False)
        :param tuple<float> edge_size: (min, max) edge size (default (15,7))
        :param tuple<float> node_size: (min, max) edge size (default (40,10))
        """
        if(node_weigth is None):
            node_weigth=[1 for k in self.classes.keys()]
        
        plt.clf()
        id_={}
        A=[]
        i=0
        G_edges=[]
        for i0,key_s in enumerate(self.classes.keys()):
            for j0,key_t in enumerate(self.classes.keys()):
                if(key_s != key_t or kwargs.get("show_internal_link", False)):
                    A.append({"Source":key_s, "Target":key_t, "Weight":0})
                    id_[(key_s, key_t)]=i
                    i+=1
                    G_edges.append((i0,j0))

        df_classe_edge=pd.DataFrame(A)
        
        N=int(len(self.edge)*min(percentage, 1))
        if(logs):
            print(f"links between classes with {percentage*100 : 0.1f}% of edges ({N} edges)")
        j=0
        for i in range(N):
            j+=1
            if(logs and j==(N//10)):
                j=0
                print(f"links between classes : {(i*100)/N : 0.1f}%")
            key_s=self.edge[self._source_col_name].iloc[i]
            key_t=self.edge[self._target_col_name].iloc[i]
            classe_s=self.node[self.node[self._id_col_name] == key_s][self._class_col_name].iloc[0]
            classe_t=self.node[self.node[self._id_col_name] == key_t][self._class_col_name].iloc[0]

            if((classe_s, classe_t) in id_.keys()):
                df_classe_edge.loc[id_[(classe_s, classe_t)], "Weight"]+=1
            
        g = igraph.Graph(directed=True)
        g.add_vertices(len(self.classes))

        for i, key in enumerate(self.classes.keys()):
            g.vs[i]["id"]= key
            g.vs[i]["label"]= "" #dictcolors[key][0]
            g.vs[i]["color"] = self.classes[key]
        
        edges_size=[df_classe_edge.loc[i, "Weight"] for i in range(len(G_edges))]
        g.add_edges(G_edges)
        
        max_edge_size,min_edge_size=kwargs.get("edge_size", (15,6)) 
         
        if(min(edges_size)!=max(edges_size)):
            g.es['width'] = [((a-min(edges_size))/(max(edges_size)-min(edges_size)))*(max_edge_size-min_edge_size)+min_edge_size for a in edges_size] #
        else:
            g.es['width'] = [(w-min(edges_size))/max(edges_size)*max_edge_size+min_edge_size for w in edges_size] #[((a-min(node_weigth))/(max(node_weigth)-min(node_weigth)))*(max_vex_size-min_vex_size)+min_vex_size for a in node_weigth] #
        

        visual_style = {}
        visual_style["bbox"] = (500,500)
        visual_style["margin"] = 50 #27

        max_node_size,min_node_size=kwargs.get("node_size", (40,10))  
        visual_style["vertex_size"] = [((a-min(node_weigth))/(max(node_weigth)-min(node_weigth)))*(max_node_size-min_node_size)+min_node_size for a in node_weigth]
        
        visual_style["vertex_label_size"] = 22
        visual_style["edge_curved"] = [0.1 for i in range(len(G_edges))]
        visual_style["edge_color"] =[self.classes[df_classe_edge.loc[i, "Source"]] for i in range(len(G_edges))]

        #igraph.plot(g, target=ax,**visual_style, edge_arrow_size= 1, directed = True)
        #plt.axis('off')
        #plt.savefig(file_name,bbox_inches = 'tight')
        out=igraph.plot(g,**visual_style, edge_arrow_size= 1, directed = True)
        out.save(file_name)

    def mostInfluentialUsersbyClass(self,file_name):
        infuence= {key:sum(self.node[self.node[self._class_col_name] == key][self._rank_col_name]) for key in self.classes.keys()}
        plt.clf() #clear graph

        plt.bar(list(self.classes.values()), [val/sum(list(infuence.values())) for _, val in infuence.items()], width=1, color=list(self.classes.values())) #sum([val for _,val in L_])
        plt.xticks(rotation=45)

        #  {sum([val for k,val in class_ if k==key])/sum([val for _,val in class_])*100 :0.1f}%
        plt.title("")
        plt.savefig(file_name, bbox_inches = 'tight')

    def makeReport(self,file_name, logs=True, **kwargs):
        """
        Create a report : uses class methods

        :param str file_name: name of the report (.docx)
        :param bool logs: if plot some logs to show the progress (default True)
        :param int nb_words_tot: number of the most common words to be displayed on the graph (kwargs default 15)
        :param int nb_words_class: number of the most common words by class to be displayed on the graph (kwargs default 10)
        :param int min_len_word: minimum size of the words to be taken into consideration (kwargs default 5)
        :param int per_edeges: must be in ]0,1], percentage of edges taken into consideration (kwargs default 1)
        :param bool show_internal_link: if true, draw an edge from the node to the same node (default False)
        :param tuple<float> edge_size: (min, max) edge size (default (15,7))
        :param tuple<float> node_size: (min, max) edge size (default (40,10))
        """
        file_class_distribution=self.out_path+f"class_distribution_{self.target_account['at']}.png"
        file_most_common_words_tot=self.out_path+f"most_common_words_tot_{self.target_account['at']}.png"
        file_most_common_words_class={key:self.out_path+f"most_common_words_{self.classes[key]}_{self.target_account['at']}.png" for key in self.classes.keys()}
        file_similarity_descriptions=self.out_path+f"similarity_descriptions_{self.target_account['at']}.png"
        file_links_between_class=self.out_path+f"links_between_class_{self.target_account['at']}.png"
        
        if(logs):
            print(f"starting {self.target_account['at']} report...")
            print(f"1st step : class distribution")
        
        pour_classe=self.class_distribution(file_class_distribution)
        
        nb_words=kwargs.get("nb_words_tot",15)
        min_len_word=kwargs.get("min_len_word",5)
        if(logs):
            print(f"2st step : most {nb_words} common words total with letters>={min_len_word}")
            
        self.most_common_words_tot(file_most_common_words_tot, nb_words=15, min_len_word=5)
        
        #############################################
        #  most common words in the description of the users by class
        #############################################
        
        nb_words_class=kwargs.get("nb_words_class",10)
        
        if(logs):
            print(f"3st step : most {nb_words_class} common words by class with letters>={min_len_word}")
            
        for key in self.classes.keys():
            rslt_df = self.node[self.node[self._class_col_name] == key]
            L_phrase=processListSentences(rslt_df.loc[:,self._desc_col_name].tolist())
            self.most_commun_words_by_class(file_most_common_words_class[key], L_phrase, nb_words_class, min_len_word=5, color=self.classes[key])
        
            #self.most_common_words_class(file_most_common_words_class[key],key_class=key, nb_words=nb_words_class, min_len_word=5)
        
        ##############################################
        # similarity of users descriptions
        #############################################
        
        if(logs):
            print(f"4st step : similarity of descriptions")
            
        self.similarity_descriptions(file_similarity_descriptions, similarityBetween)
        
        if(logs):
            print(f"5st step : links between class /!\ It takes time !")
            
        self.links_between_class(file_links_between_class, percentage=kwargs.get("per_edeges",1), 
                                    node_weigth=pour_classe, logs=logs, 
                                    show_internal_link=kwargs.get("show_internal_link", False),
                                    node_size=kwargs.get("node_size", (40,10)),
                                    edge_size=kwargs.get("edge_size", (15,7)))
        
        most_influential_accounts=[]
        for key in self.classes.keys():
            rslt_df = self.node[self.node[self._class_col_name] == key].sort_values(by=[self._rank_col_name], ascending=False)
            R=[]
            for i in range(min(len(rslt_df),10)):
                R.append({"name":str(rslt_df[self._name_col_name].iloc[i]),
                          "at":"@"+str(rslt_df[self._at_col_name].iloc[i]),
                          "description":str(rslt_df[self._desc_col_name].iloc[i])
                         })
            most_influential_accounts.append(R)
        
        ###
        # on somme les pages rank et on voit le résultat
        ##
        file_mostInfluentialUsersbyClass=self.out_path+f"file_mostInfluentialUsersbyClass_{self.target_account['at']}.png"
        self.mostInfluentialUsersbyClass(file_mostInfluentialUsersbyClass)
        
        ################
        # Tweet analyse
        ################
        content_class=[]
        for i, (key,color) in enumerate(self.classes.items()): 
            L={"title":{f"Class {color} ({pour_classe[i]*100 : .1f}%)"},
               "text_1":""}
            K=[{"title":f"Most {nb_words_class} common words",
               "text_1":f"We only check for words with more than {min_len_word} characters",
                "img":file_most_common_words_class[key]},
                {"title":"Most influential accounts",
                "text_1":[f"{acc['name']} ({acc['at']}) : «{acc['description']}» " for acc in most_influential_accounts[i]]
                }
            ]
            if(self._tweets):
                #most_common_words in tweet by _class 
                rslt_df = self._df_tweets[self._df_tweets[self._class_col_name] == key]
                L_phrase=processListSentences(rslt_df.loc[:,"text"].tolist())
                self.most_commun_words_by_class(self.out_path+f"tweet_{key}.png", L_phrase, 10, min_len_word=5, color=color, title=f"The {10} most common word in tweets of {color} class")
                
                mini=min([datetime.datetime.fromisoformat(d) for d in rslt_df.loc[:,"created_at"]])
                maxi=max([datetime.datetime.fromisoformat(d) for d in rslt_df.loc[:,"created_at"]])

                K.append({"title":"Most commun words in tweets",
                          "text_1":{"text":f"{len(L_phrase)} tweets between {mini} and {maxi} UTC", "type":["italic"]},
                          "img":self.out_path+f"tweet_{key}.png"})

                ##3 comptes ayant tweeté le plus
                users_t=collections.Counter(rslt_df.loc[:,"user_id"].tolist()).most_common(3)
                K.append({"title":"Users who tweeted the most",
                          "text_1":[f"""@{self.node[self.node['Id'] == int(user_id)]['screen_name'].tolist()[0] if not self.node[self.node['Id'] == int(user_id)].empty else None} «{self.node[self.node['Id'] == int(user_id)]['description'].tolist()[0] if not self.node[self.node['Id'] == int(user_id)].empty else None}» tweet {time} time""" for user_id,time in users_t]})

            content_class.append([L, K])

        content=[
            {"title":f'Analysis of {self.target_account["name"]} ({self.target_account["at"]})\'s Graph',"text_1":''},
            [
                {"title":"Introduction", "text_1":{"text":"<insert a comment about #Bucha>", "type":["italic"]}, "text_2":"On the document below, we will first see a global analysis, and a comparative analysis of the class, then we will see in each class the most influential accounts"},
                [
                    {"title":"Class distribution", 
                    "text_1":f"All individuals are placed into one of {len(self.classes)} class :",
                    "img":file_class_distribution},
                    {"title":"Links between class", 
                    "text_1":f"The most interesting analysis for the comparison between the class, is to show links between class.",
                    "img":file_links_between_class,
                    "text_2":{"text":"All accounts of a class are merged to a sigle node and we see that the links that leave the class \n <insert a comment about this graph>", "type":["italic"]}},
                    {"title":"Most common words in descriptions", 
                    "text_1":f"For an initial understanding, the target's graph, we can see the {nb_words} most commun words on follower's descriptions spread on each classe",
                    "img":file_most_common_words_tot},
                    {"title":"Similarity of descriptions", 
                    "text_1":f"Similarity between the follower's description and the description of {self.target_account['name']} \n {self.target_account['description']}",
                    "img":file_similarity_descriptions,
                    "text_2":f"Warning: This methode use only the similarity between {self.target_account['name']}'s description and follower's description, and not the meaning of the description"}
                ],
                {"title":"Class analysis","text_1":""},
                content_class,
                [{"title":"other analyses","text_1":""},
                    [{"title":"Distribution of the most influential users by class",
                    "text_1":"With the sum of the page rank, we can estimate the influence of the classes.",
                    "img":file_mostInfluentialUsersbyClass},
                    ]
                ]
            ]
        ]
        if(self._tweets):
            self.pieGraph(self.out_path+f"tweet_interaction.png", collections.Counter(self._df_tweets["interaction"].tolist()), "type of tweet interaction")
            content[-1][-1][-1].append({"title":"type of tweet interaction",
                    "text_1":"type of tweet interaction",
                    "img":self.out_path+f"tweet_interaction.png"
            })

            self.pieGraph(self.out_path+f"tweet_langage.png", collections.Counter(self._df_tweets["language"].tolist()), "language of the tweets")
            content[-1][-1][-1].append({"title":"language of the tweets",
                    "text_1":"language of the tweets",
                    "img":self.out_path+f"tweet_langage.png"
            })

        writeDoc(file_name, content)
    
if __name__ == '__main__':
    r=autoReport("nodes.csv", "edges.csv", {2339:"magenta", 2433:"limegreen", 2566:"deepskyblue", 1052:"saddlebrown", 2498:"darkorange", 1583:"red"}, 
           {"name":"Antoine Bondaz", "at":"@AntoineBondaz", "description":"""Foodie - 🕵🏼‍Research @FRS_org - 👨🏼‍🏫 Teach @SciencesPo - 🇨🇳🇹🇼🇰🇷🇰🇵's foreign & security policies - Ph.D."""})

    r.makeReport("out.docx")